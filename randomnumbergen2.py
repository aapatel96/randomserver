##random number generator server

import random, string, pymongo, telegram, os, codecs, boto3, re, os
from docx import Document
from docx import Document
from flask import Flask, request, jsonify
import sys 

reload(sys)  
sys.setdefaultencoding('utf8')

bot = telegram.Bot('395089971:AAGdmNnerGxByQZqhjen2hAGIZ2CBW-WcnY')

s3 = boto3.client(
    's3',
    aws_access_key_id = os.environ['AWS_ACCESS_KEY_ID'],
    aws_secret_access_key = os.environ['AWS_SECRET_ACCESS_KEY']
)

app = Flask(__name__)

@app.route("/create", methods=['POST'])
def randomword():
   name = request.form['name']
   address = request.form['address']
   passport_number = request.form['passport_number']
   poatemplate = open('template2.txt','r')
   poatext = poatemplate.read()
   poatext = poatext.decode('unicode_escape').encode('utf-8')

   poatext= poatext.replace('GRANTORNAME',name)
   poatext=poatext.replace('GRANTORADDRESS',address)
   poatext=poatext.replace('GRANTORPASSPORTNUMBER',passport_number)

   userfilename = name+'-poa.txt'
   userfile = open(userfilename,'w')
   userfile.write(poatext)
   userfile.close()  
   userfiletext = open(userfilename,'r').read()
   userfiletext = re.sub(r'[^\x00-\x7F]+|\x0c',' ', userfiletext)

   file_text_comps = userfiletext.split('[paragraph]')
   
   document = Document()
   document.add_heading(file_text_comps[0], level=1)
   file_text_comps= file_text_comps[1:]
   count =0
   for paragraph in file_text_comps:
      count = count+ 1
      print count
      document.add_paragraph(paragraph)
   docname= name+'.docx'
   document.save(name+'.docx')
   s3.upload_file(docname,'powerofattorneybot',docname)
   userfilenamealphabets = []
   for char in docname:
      if char == ' ':
         userfilenamealphabets.append('+')
      else:
         userfilenamealphabets.append(char)

   aws_url_filename= ''.join(userfilenamealphabets)
   aws_url='https://s3-us-west-1.amazonaws.com/powerofattorneybot/'+aws_url_filename
   
##   bot.send_document(chat_id='89380112',document=userfile)

   try:
      return jsonify(status='success',name=name,address = address,passport_number=passport_number,url=aws_url,paras=str(len(file_text_comps)))
   except:
      return jsonify(status='failed',error="Please pass numbers only")

##   bot.send_message(chat_id='89380112',text=length)
##   bot.send_message(chat_id='89380112',text=''.join(random.choice(string.lowercase) for i in range(length)))
##   return str(length)
##   return ''.join(random.choice(string.lowercase) for i in range(length))

if __name__ == "__main__":
    PORT = int(os.environ.get('PORT', '5000'))

    app.run(
        host="0.0.0.0",
        port=PORT
    )